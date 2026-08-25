from __future__ import annotations

import csv
import json
import math
import os
from collections import Counter
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\repos\biofin-research\budget_biodiv_cls3")
REPORT_DIR = ROOT / "reports"
ASSET_DIR = REPORT_DIR / "report_assets"
OUT = REPORT_DIR / "BIOFIN_예산사업_분류모형_개발_결과_비교평가_보고서_수정본.docx"

L1 = ROOT / "llm/v1/outputs/260812_2023data"
L2 = ROOT / "llm/v2/outputs/260818_2023data"
T1 = ROOT / "transformer/v1/outputs/260812_category_v2"
T2 = ROOT / "transformer/v2/outputs/260818_2023data"

NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "DCE6F1"
LIGHT = "F2F4F7"
MID = "D9E2F3"
DARK = "1F1F1F"
MUTED = "666666"
RED = "9C0006"
GOLD = "806000"
WHITE = "FFFFFF"


def read_json(path: Path):
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def read_csv(path: Path):
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def safe_float(v):
    try:
        return float(v)
    except (TypeError, ValueError):
        return math.nan


def pct(v, digits=1):
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return "-"
    return f"{v * 100:.{digits}f}%"


def fmt(v, digits=3):
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return "-"
    return f"{v:.{digits}f}"


def calc_subset(rows, true_col, pred_col, exclude_zero=False):
    pairs = []
    for row in rows:
        t = str(row.get(true_col, "")).strip()
        p = str(row.get(pred_col, "")).strip()
        if not t or not p:
            continue
        if exclude_zero and t == "0":
            continue
        pairs.append((t, p))
    return {
        "n": len(pairs),
        "correct": sum(t == p for t, p in pairs),
        "accuracy": sum(t == p for t, p in pairs) / len(pairs) if pairs else math.nan,
    }


def binary_metrics(rows, true_col, pred_col):
    pairs = []
    for row in rows:
        t = str(row.get(true_col, "")).strip()
        p = str(row.get(pred_col, "")).strip()
        if not t or not p:
            continue
        pairs.append((t != "0", p != "0"))
    tp = sum(t and p for t, p in pairs)
    fp = sum((not t) and p for t, p in pairs)
    fn = sum(t and (not p) for t, p in pairs)
    tn = sum((not t) and (not p) for t, p in pairs)
    precision = tp / (tp + fp) if tp + fp else math.nan
    recall = tp / (tp + fn) if tp + fn else math.nan
    f1 = 2 * precision * recall / (precision + recall) if precision + recall else math.nan
    return {"tp": tp, "fp": fp, "fn": fn, "tn": tn, "precision": precision, "recall": recall, "f1": f1}


def label_sort(x):
    try:
        return tuple(int(p) for p in str(x).split("."))
    except ValueError:
        return (999, str(x))


UPPER_NAMES = {
    "0": "비해당",
    "1": "유전자원 접근 및 이익공유(ABS)",
    "2": "생물다양성 인식 제고 및 연구",
    "3": "생물안전성",
    "4": "친환경 경제 전환",
    "5": "생물다양성 기획 및 재정",
    "6": "오염 저감",
    "7": "보호지역 및 보전조치(PA & OECM)",
    "8": "생태계 복원",
    "9": "지속가능한 자연 이용",
}

UPPER_SUMMARIES = {
    "0": "직접·간접 생물다양성 활동이 확인되지 않는 일반 행정·운영·소득지원 등",
    "1": "유전자원 접근, 나고야의정서 이행, 유전자원 정보·이익공유 기반",
    "2": "생물다양성 교육·인식·조사·연구·모니터링 및 관련 과학기술 개발",
    "3": "침입외래종·병해충 유입 차단, 검역·예찰·방제 및 GMO/LMO 안전관리",
    "4": "저탄소·순환경제, 지속가능 에너지·관광·도시와 녹색 공급망 전환",
    "5": "생물다양성 법률·정책·재정·조정, 환경평가·공간계획·국제협약",
    "6": "토양·수질·하수·폐기물·해양쓰레기 등 오염의 예방·저감·관리",
    "7": "보호지역·OECM 관리·확장, 완충지와 야생종·이동성 종 보전",
    "8": "훼손된 생태계·서식지의 복원·재활과 복원지 사후관리",
    "9": "농업·양식·어업·임업·토지·연안의 지속가능한 이용과 관리",
}

SUB_NAMES = {
    "0": "비해당", "1.01": "유전자원 스크리닝·접근 허가", "1.04": "나고야의정서 이행",
    "1.05": "유전자원 정보 접근권", "2.01": "정규 교육", "2.02": "비정규 교육·기술훈련",
    "2.03": "대중 인식·소통", "2.04": "생물다양성 과학 연구", "2.06": "CBD 정보공유체계",
    "3.01": "침입외래종 관리", "3.02": "GMO/LMO 관리", "4.01": "녹색 공급망",
    "4.04": "지속가능 에너지", "4.05": "지속가능 관광", "4.07": "지속가능 도시·지역",
    "5.01": "생물다양성 법률·정책", "5.02": "타 부처·섹터 법률·정책", "5.03": "부처 간 조정·관리",
    "5.04": "생물다양성 재정", "5.05": "전략환경평가 체계", "5.06": "공간계획", "5.07": "다자간 환경협약",
    "6.01": "토양·수질 오염", "6.03": "하수·폐기물", "6.04": "연안·해양 쓰레기",
    "6.05": "기타 오염", "6.06": "오염영향 관리 기반", "7.01": "보호지역·ICCA 관리",
    "7.02": "보호구역 외부 완충지", "7.03": "OECM", "7.04": "야생종·이동성 종 보전",
    "8.02": "훼손지 공학적 복원", "8.03": "복원지 사후관리", "9.01": "농업생물다양성",
    "9.02": "지속가능 농업", "9.03": "지속가능 양식", "9.04": "지속가능 어업",
    "9.05": "지속가능 임업", "9.06": "지속가능 토지관리", "9.07": "지속가능 연안·해양",
}


def load_data():
    d = {
        "l1_sum": read_json(L1 / "run_summary.json"),
        "l1_met": read_json(L1 / "evaluation_metrics.json"),
        "l2_sum": read_json(L2 / "run_summary.json"),
        "l2_met": read_json(L2 / "evaluation_metrics.json"),
        "t1_split": read_json(T1 / "split_summary.json"),
        "t2_split": read_json(T2 / "split_summary.json"),
        "t1_met": read_json(T1 / "metrics.json"),
        "t2_met": read_json(T2 / "metrics.json"),
        "t1_train": read_csv(T1 / "train_log.csv"),
        "t2_train": read_csv(T2 / "train_log.csv"),
        "t1_under": read_json(T1 / "train_undersampling_summary.json"),
        "t2_under": read_json(T2 / "train_undersampling_summary.json"),
        "t1_pred": read_csv(T1 / "test_predictions.csv"),
        "t2_pred": read_csv(T2 / "test_predictions.csv"),
        "t1_dist": read_csv(T1 / "label_distribution.csv"),
        "t2_dist": read_csv(T2 / "label_distribution.csv"),
        "l1_rows": read_csv(L1 / "260812_2023data_llm_classified.csv"),
        "l2_rows": read_csv(L2 / "260812_2023data_llm_classified.csv"),
        "t1_report": parse_report(T1 / "classification_report.txt"),
        "t2_report": parse_report(T2 / "classification_report.txt"),
        "l1_cm": read_csv(L1 / "confusion_matrix.csv"),
        "l2_cm": read_csv(L2 / "confusion_matrix.csv"),
        "t1_cm": read_csv(T1 / "confusion_matrix.csv"),
        "t2_cm": read_csv(T2 / "confusion_matrix.csv"),
    }
    d["l1_nz"] = calc_subset(d["l1_rows"], "BIOFIN 1차 카테고리", "LLM BIOFIN 1차 카테고리", True)
    official_nonzero = sum(x["support"] for code, x in d["l1_met"]["per_class"].items() if code != "0")
    d["l1_official_nz"] = {
        "n": official_nonzero,
        "correct": d["l1_met"]["correct_rows"],
        "accuracy": d["l1_met"]["correct_rows"] / official_nonzero,
    }
    d["l1_reconstructed"] = calc_subset(d["l1_rows"], "BIOFIN 1차 카테고리", "LLM BIOFIN 1차 카테고리", False)
    d["l2_nz"] = calc_subset(d["l2_rows"], "하위 카테고리", "LLM BIOFIN 하위 카테고리", True)
    d["t1_nz"] = calc_subset(d["t1_pred"], "true_label", "pred_label", True)
    d["t2_nz"] = calc_subset(d["t2_pred"], "true_subcategory", "pred_subcategory", True)
    d["l1_bin"] = binary_metrics(d["l1_rows"], "BIOFIN 1차 카테고리", "LLM BIOFIN 1차 카테고리")
    d["l2_bin"] = binary_metrics(d["l2_rows"], "하위 카테고리", "LLM BIOFIN 하위 카테고리")
    d["t1_bin"] = binary_metrics(d["t1_pred"], "true_label", "pred_label")
    d["t2_bin"] = binary_metrics(d["t2_pred"], "true_subcategory", "pred_subcategory")
    return d


def parse_report(path: Path):
    out = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        parts = line.split()
        if len(parts) == 5 and parts[0] not in {"accuracy", "macro", "weighted"}:
            try:
                out[parts[0]] = {"precision": float(parts[1]), "recall": float(parts[2]), "f1": float(parts[3]), "support": int(parts[4])}
            except ValueError:
                pass
    return out


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths=None, font_size=8.5, header_fill=MID):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        set_cell_shading(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(NAVY)
            r.font.size = Pt(font_size)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.size = Pt(font_size)
        if len(table.rows) % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc, title, text, tone="blue"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    fill = "EAF2F8" if tone == "blue" else "FFF2CC"
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY if tone == "blue" else GOLD)
    p.add_run(text)
    set_table_widths(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def page_break(doc):
    doc.add_page_break()


def set_repeat_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        s = styles[name]
        s.font.name = "Malgun Gothic"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    styles["Caption"].font.name = "Malgun Gothic"
    styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Caption"].font.size = Pt(9)
    styles["List Bullet"].font.name = "Malgun Gothic"
    styles["List Bullet"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "BIOFIN 예산사업 분류모형 개발 결과 및 비교·평가"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs:
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer
    set_repeat_page_number(footer.paragraphs[0])


def make_charts(d):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.family"] = "Malgun Gothic"
    plt.rcParams["axes.unicode_minus"] = False
    models = ["LLM v1", "Transformer v1", "LLM v2", "Transformer v2"]
    accuracy = [d["l1_met"]["accuracy"], d["t1_met"]["test"]["accuracy"], d["l2_met"]["accuracy"], d["t2_met"]["test"]["accuracy"]]
    macro = [d["l1_met"]["macro_f1"], d["t1_met"]["test"]["macro_f1"], d["l2_met"]["macro_f1"], d["t2_met"]["test"]["macro_f1"]]
    x = range(4)
    fig, ax = plt.subplots(figsize=(8.4, 4.2))
    ax.bar([i - .18 for i in x], accuracy, .36, label="Accuracy", color="#2E74B5")
    ax.bar([i + .18 for i in x], macro, .36, label="Macro F1", color="#9EADBA")
    ax.set_xticks(list(x), models)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("점수")
    ax.set_title("모형별 보고 성능 (평가 모집단이 서로 다름)")
    ax.legend(frameon=False, ncol=2, loc="upper left")
    ax.grid(axis="y", alpha=.2)
    for i, v in enumerate(accuracy): ax.text(i-.18, v+.025, f"{v:.3f}", ha="center", fontsize=8)
    for i, v in enumerate(macro): ax.text(i+.18, v+.025, f"{v:.3f}", ha="center", fontsize=8)
    fig.tight_layout()
    p = ASSET_DIR / "overall_metrics.png"
    fig.savefig(p, dpi=180, bbox_inches="tight")
    plt.close(fig)

    labels = [r["label"] for r in d["t1_dist"]]
    vals = [int(r["count"]) for r in d["t1_dist"]]
    fig, ax = plt.subplots(figsize=(8.4, 4.0))
    colors = ["#2E74B5" if l == "0" else "#9EADBA" for l in labels]
    ax.bar(labels, vals, color=colors)
    ax.set_yscale("log")
    ax.set_xlabel("1차 카테고리")
    ax.set_ylabel("건수(로그 척도)")
    ax.set_title("1차 카테고리 전체 분포: 클래스 0의 압도적 우세")
    ax.grid(axis="y", alpha=.2)
    fig.tight_layout()
    p2 = ASSET_DIR / "upper_distribution.png"
    fig.savefig(p2, dpi=180, bbox_inches="tight")
    plt.close(fig)

    fig, axes = plt.subplots(1, 2, figsize=(8.4, 3.7), sharey=True)
    for ax, rows, title in [(axes[0], d["t1_train"], "Transformer v1"), (axes[1], d["t2_train"], "Transformer v2")]:
        ep = [int(r["epoch"]) for r in rows]
        ma = [safe_float(r["macro_f1"]) for r in rows]
        ac = [safe_float(r["accuracy"]) for r in rows]
        ax.plot(ep, ma, marker="o", label="Macro F1", color="#2E74B5")
        ax.plot(ep, ac, marker="s", label="Accuracy", color="#7F8C8D")
        ax.set_xticks(ep)
        ax.set_title(title)
        ax.set_xlabel("Epoch")
        ax.grid(alpha=.2)
    axes[0].set_ylabel("검증 점수")
    axes[1].legend(frameon=False, fontsize=8)
    fig.suptitle("검증 성능 추이")
    fig.tight_layout()
    p3 = ASSET_DIR / "train_curves.png"
    fig.savefig(p3, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return p, p2, p3


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    shape = r.add_picture(str(path), width=Inches(6.25))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", Path(path).stem)
    p.paragraph_format.space_after = Pt(2)
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    for run in cp.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def build():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    d = load_data()
    charts = make_charts(d)
    doc = Document()
    setup_styles(doc)

    # Cover
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("연구 결과 보고서")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(BLUE)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BIOFIN 예산사업 분류모형 개발 결과\n및 비교·평가 보고서")
    r.bold = True; r.font.size = Pt(25); r.font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph("LLM·Transformer 기반 1차 및 하위 카테고리 분류모형")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(MUTED)
    for _ in range(7): doc.add_paragraph()
    p = doc.add_paragraph("분석자료  2023년 예산사업 3,972건\n작성일  2026년 8월")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(DARK)
    page_break(doc)

    # Static table of contents: stable across Word/LibreOffice without field refresh.
    add_heading(doc, "목차", 1)
    toc_items = [
        ("요약", "분석 목적, 주요 성능, 활용 권고"),
        ("제1장. 연구 개요", "배경, 목적, 분석 대상 및 비교 원칙"),
        ("제2장. 분류체계 및 원천데이터", "BIOFIN 체계, 입력자료, 전처리, 불균형"),
        ("제3장. 학습·검증·테스트 데이터 구성", "그룹 분할, 클래스별 구성, Coverage"),
        ("제4장. 모형 구축 및 평가 방법", "LLM, Transformer, 평가 지표"),
        ("제5장. 1차 카테고리 분류 결과", "LLM v1, Transformer v1 및 비교"),
        ("제6장. 하위 카테고리 분류 결과", "LLM v2, Transformer v2 및 비교"),
        ("제7장. 네 모형 종합 비교", "0 포함·제외, 이진 탐지, 오류 및 운영 특성"),
        ("제8장. 결과 해석과 한계", "평가·레이블 한계"),
        ("제9장. 개선 방안 및 활용 제안", "개선 우선순위와 반자동 운영"),
        ("제10장. 결론", "핵심 결과와 후속 과제"),
        ("부록", "코드표, 분포, 산출물 및 재현 설정"),
    ]
    add_table(doc, ["구분", "주요 내용"], toc_items, [3000, 6360], 9.0, header_fill=LIGHT)
    page_break(doc)

    # Executive summary
    add_heading(doc, "요약", 1)
    add_callout(doc, "핵심 결론", "현 산출물만으로 완전 자동 확정을 권고하기는 어렵다. 1차 분류는 Transformer v1이 LLM v1보다 안정적이지만 소수 클래스 성능이 제한적이고, 하위 분류는 LLM v2가 높은 전체 정확도를 보이나 클래스 0의 영향이 크다. 공통 홀드아웃 테스트 세트와 사람 검토를 결합한 반자동 운영이 적절하다.")
    add_heading(doc, "1. 분석 목적 및 범위", 2)
    add_para(doc, "본 분석은 2023년 예산사업 3,972건을 BIOFIN 분류체계에 따라 자동 분류하기 위해 개발된 네 가지 모형의 산출물을 정리하고, 성능·오류·운영 가능성을 비교하는 데 목적이 있다. 비교 대상은 1차 카테고리용 LLM v1과 Transformer v1, 하위 카테고리용 LLM v2와 Transformer v2이다.")
    add_bullets(doc, [
        "1차 카테고리 분류: 비해당(0)과 9개 BIOFIN 상위 범주를 구분한다.",
        "하위 카테고리 분류: 비해당(0)과 실제 학습자료에 존재하는 39개 하위 코드를 구분한다.",
        "LLM은 전체 입력 중 유효 예측이 저장된 행을 평가했고, Transformer는 그룹 분할한 독립 테스트 397건을 평가했다.",
        "평가 모집단이 다르므로 네 모형의 숫자를 단순 순위로 해석하지 않고 동일 분류 수준 내에서 비교했다.",
    ])
    add_heading(doc, "2. 주요 성능 비교", 2)
    rows = [
        ["LLM v1", "1차(10)", "2,905 / 3,972", pct(d["l1_met"]["accuracy"]), fmt(d["l1_met"]["macro_f1"]), pct(d["l1_met"]["evaluated_rows"] / 3972), "0 예측 저장·평가 결함"],
        ["Transformer v1", "1차(10)", "테스트 397", pct(d["t1_met"]["test"]["accuracy"]), fmt(d["t1_met"]["test"]["macro_f1"]), "100.0%", "독립 테스트"],
        ["LLM v2", "하위(51 후보·43 활성)", "3,430 / 3,972", pct(d["l2_met"]["accuracy"]), fmt(d["l2_met"]["macro_f1"]), pct(d["l2_met"]["evaluated_rows"] / 3972), "클래스 0 영향 큼"],
        ["Transformer v2", "하위(40 학습)", "테스트 397", pct(d["t2_met"]["test"]["accuracy"]), fmt(d["t2_met"]["test"]["macro_f1"]), "100.0%", "독립 테스트"],
    ]
    add_caption(doc, "표 1. 네 모형의 핵심 성능")
    add_table(doc, ["모형", "분류 수준", "평가 건수", "Accuracy", "Macro F1", "Coverage", "해석상 주의"], rows, [1200, 1200, 1400, 1050, 1050, 1050, 2405], 7.8)
    add_para(doc, "주: LLM v2의 51은 평가 코드 목록 전체(0 포함), 43은 정답 또는 예측이 한 건 이상 존재한 활성 코드 수이다. Transformer v2의 40은 학습자료에 실제 존재하여 모델 출력층에 포함된 코드 수로, 서로 다른 개념이다.", italic=True, color=MUTED)
    add_picture(doc, charts[0], "그림 1. 산출물에 기록된 Accuracy와 Macro F1. 평가 모집단 차이 때문에 모형 간 직접 순위로 사용할 수 없다.")
    add_bullets(doc, [
        "LLM v1: 평가상 정확도 11.60%, Macro F1 0.271이다. 다만 캐시에는 0 예측 1,062건이 있으나 평가에서는 0 예측이 0건으로 집계되어 결과가 파이프라인 오류의 영향을 받았다.",
        "Transformer v1: 테스트 정확도 60.45%, Macro F1 0.233이다. 클래스 0 재현율은 58.43%이고, 클래스 2·6·9는 재현율이 높지만 정밀도가 낮다.",
        "LLM v2: 평가 정확도 95.36%이나 비해당 0이 평가자료의 89.48%(3,069/3,430)를 차지한다. 하위 관련 클래스만 보면 정확도는 별도로 낮아진다.",
        "Transformer v2: 테스트 정확도 61.46%, Macro F1 0.225이다. 테스트 397건 중 클래스 0이 337건으로, 전체 정확도 역시 다수 클래스 성능의 영향을 받는다.",
    ])
    add_heading(doc, "3. 최종 활용 권고", 2)
    add_bullets(doc, [
        "단기: LLM v2 또는 Transformer 결과를 1차 후보로 제시하고, 신뢰도·클래스별 위험도에 따라 사람이 확정한다.",
        "중기: 동일한 공통 테스트 세트에서 LLM과 Transformer를 재평가하고, 클래스 0 포함·제외 및 이진 탐지 지표를 함께 관리한다.",
        "장기: 1차 카테고리 판정 후 하위 카테고리를 분류하는 계층형 구조와 두 모형의 합의 기반 앙상블을 구축한다.",
    ])

    page_break(doc)
    add_heading(doc, "제1장. 연구 개요", 1)
    add_heading(doc, "1. 연구 배경", 2)
    add_para(doc, "BIOFIN 예산분류는 정부 예산사업 중 생물다양성에 직접 또는 간접적으로 기여하는 활동을 식별하고 재정 흐름을 파악하기 위한 기초 작업이다. 사업명과 예산정보만으로는 목적과 활동을 충분히 구분하기 어려우므로 사업설명자료 본문을 함께 검토해야 한다. 수천 건을 사람이 일관된 기준으로 판정하는 데에는 많은 시간과 전문성이 요구되며, 경계 사례에서는 판정자 간 차이도 발생할 수 있다.")
    add_para(doc, "본 연구는 생성형 언어모형과 지도학습 Transformer를 병행하여 자동분류의 가능성을 검토했다. LLM은 상세 기준을 프롬프트로 제공하면 별도 학습 없이 근거 문장을 생성할 수 있고, Transformer는 정답 자료를 통해 반복 가능한 분류경계를 학습한다. 두 접근은 상호 대체재라기보다 정확도·설명력·운영비용을 보완하는 후보로 볼 수 있다.")
    add_heading(doc, "2. 연구 목적", 2)
    add_bullets(doc, ["1차 및 하위 카테고리 자동분류 가능성 평가", "LLM과 Transformer의 정량·정성 성능 비교", "클래스별 취약점과 데이터 불균형 영향 진단", "실무 적용을 위한 검토 절차와 개선과제 제안"])
    add_heading(doc, "3. 분석 대상 모형", 2)
    model_rows = [
        ["LLM v1", "gemma3:12b", "1차 카테고리 0~9", "메타데이터+본문 최대 16,000자", d["l1_sum"]["prompt_version"]],
        ["LLM v2", "gemma3:12b", "0+50개 후보(43개 활성)", "메타데이터+본문 최대 16,000자", d["l2_sum"]["prompt_version"]],
        ["Transformer v1", "KLUE/RoBERTa-base", "1차 카테고리 0~9", "본문 512토큰 chunk+attention", "학습형"],
        ["Transformer v2", "KLUE/RoBERTa-base", "0+39개 하위 코드", "본문 512토큰 chunk+attention", "학습형"],
    ]
    add_caption(doc, "표 2. 분석 대상 모형")
    add_table(doc, ["모형", "기반 모델", "출력", "입력·구조", "버전"], model_rows, [1100, 1450, 1500, 2600, 2710], 8)
    add_heading(doc, "4. 비교 원칙", 2)
    add_bullets(doc, ["동일 분류 수준끼리 비교한다.", "LLM의 유효예측 평가와 Transformer의 독립 테스트 평가를 구분한다.", "Accuracy와 함께 Macro F1, 클래스별 Recall, Support를 본다.", "클래스 0 포함 결과와 관련 클래스만의 결과를 분리한다.", "Support가 매우 작은 클래스는 수치보다 표본 부족 자체를 핵심 결과로 해석한다."])

    add_heading(doc, "제2장. 분류체계 및 원천데이터", 1)
    add_heading(doc, "1. BIOFIN 분류체계", 2)
    add_para(doc, "분류체계는 비해당 0과 생물다양성 재정활동 9개 상위 범주로 구성된다. 하위 분류는 활동 유형을 더 세분화한다. LLM v2 평가는 전체 51개 후보 코드를 정의했으며 그중 정답 또는 예측이 존재한 활성 코드는 43개였다. Transformer v2는 학습자료에 실제 존재한 비해당 포함 40개 코드만 출력층으로 구성했다.")
    upper_rows = [[k, v, UPPER_SUMMARIES[k]] for k, v in UPPER_NAMES.items()]
    add_caption(doc, "표 3. BIOFIN 1차 카테고리")
    add_table(doc, ["코드", "명칭", "판정 개요"], upper_rows, [700, 3500, 5160], 8.5)
    add_heading(doc, "2. 원천데이터 및 입력 변수", 2)
    add_para(doc, "원천자료는 2023년 예산사업 3,972건이다. 주요 필드는 회계연도, 소관명, 회계·계정, 분야·부문, 프로그램, 단위사업, 세부사업, 예산액, 기존 BIOFIN 레이블과 사업설명자료 경로이다. LLM은 이 메타데이터와 본문을 함께 사용했고, Transformer는 document_only 설정에 따라 사업설명자료 본문 추출 성공 건을 학습에 사용했다.")
    add_heading(doc, "3. 전처리 및 문서 매칭", 2)
    add_bullets(doc, [
        "연도·소관·세부사업명 등의 정보를 정규화하여 예산행과 사업설명자료를 매칭했다.",
        "Transformer 산출물의 dataset_match_success는 3,972건이며, 원천 문서 전체 중 레이블 데이터에 포함되지 않은 문서는 별도 실패 로그에 기록되었다.",
        "동일 소관·세부사업은 회계연도를 제외한 그룹 키로 묶어 서로 다른 분할에 섞이지 않도록 했다.",
        "하위 코드의 0.0, 2.040과 같은 표기를 0, 2.04로 정규화하고 내부 정수 ID로 인코딩했다.",
    ])
    add_heading(doc, "4. 클래스 분포와 불균형", 2)
    total0 = int(next(r["count"] for r in d["t1_dist"] if r["label"] == "0"))
    add_para(doc, f"1차 카테고리에서 클래스 0은 {total0:,}건으로 전체 3,972건의 {total0/3972:.1%}를 차지한다. 나머지 클래스는 8~137건에 불과하다. 하위 분류에서는 희소성이 더 심해 다수 코드가 전체 자료에서도 한 자릿수 표본만 가진다.")
    add_picture(doc, charts[1], "그림 2. 1차 카테고리별 전체 건수(로그 척도)")
    add_callout(doc, "해석 원칙", "불균형 자료에서는 0만 잘 맞혀도 Accuracy가 높아질 수 있다. Macro F1은 각 클래스를 같은 비중으로 평균하므로 소수 클래스 성능을 더 직접적으로 보여준다.", "yellow")

    page_break(doc)
    add_heading(doc, "제3장. 학습·검증·테스트 데이터 구성", 1)
    add_heading(doc, "1. 데이터 분할 원칙", 2)
    add_para(doc, "Transformer는 소관명과 세부사업명을 기준으로 동일 사업 그룹을 보존하면서 8:1:1로 분할했다. 두 실험 모두 학습 3,178건, 검증 397건, 테스트 397건이며, 그룹 중복과 상충 레이블 그룹은 0건으로 기록되었다. 이 구조는 동일 사업 문서가 학습과 테스트에 동시에 포함되는 누수를 방지한다.")
    split_rows = [
        ["Transformer v1", "3,178", "397", "397", d["t1_split"]["train_groups"], d["t1_split"]["overlap_groups"], "1차 10개"],
        ["Transformer v2", "3,178", "397", "397", d["t2_split"]["train_groups"], d["t2_split"]["overlap_groups"], "하위 40개"],
    ]
    add_caption(doc, "표 4. Transformer 데이터 분할")
    add_table(doc, ["모형", "학습", "검증", "테스트", "학습 그룹", "중복 그룹", "레이블"], split_rows, [1400, 900, 900, 900, 1200, 1200, 2860], 8.5)
    add_para(doc, "주: 학습행 3,178건과 학습 그룹 3,168개의 차이 10건은 오타가 아니다. 학습 세트에는 동일 소관·세부사업 그룹에 2개 행이 속한 그룹 5개, 3개 행 1개, 4개 행 1개가 있어 7개 그룹에서 총 10개의 추가 행이 발생했다.", italic=True, color=MUTED)
    add_heading(doc, "2. Transformer v1 데이터", 2)
    v1rows=[]
    for label in sorted(d["t1_split"]["train_label_counts"], key=label_sort):
        v1rows.append([label, UPPER_NAMES.get(label,""), d["t1_split"]["train_label_counts"].get(label,0), d["t1_split"]["valid_label_counts"].get(label,0), d["t1_split"]["test_label_counts"].get(label,0)])
    add_caption(doc, "표 5. Transformer v1 클래스별 분할")
    add_table(doc, ["코드", "명칭", "학습", "검증", "테스트"], v1rows, [650, 3950, 1200, 1200, 2360], 8)
    u=d["t1_under"]
    add_para(doc, f"불균형 완화를 위해 학습 세트의 클래스 0을 {u['majority_before']:,}건에서 {u['majority_after']:,}건으로 축소했다. 이에 따라 실제 학습 투입 건수는 {u['train_before']:,}건에서 {u['train_after']:,}건으로 줄었고, 검증·테스트 분포는 원래 상태를 유지했다.")
    add_heading(doc, "3. Transformer v2 데이터", 2)
    add_para(doc, "하위 분류는 40개 클래스이지만 검증·테스트 각각에 모든 클래스가 존재하지 않는다. 테스트에서는 25개 코드만 정답 표본을 가지며, 관련 클래스 60건 중 상당수가 클래스당 1~3건이다. 개별 클래스의 0% 또는 100% 성능은 일반화 성능이라기보다 소수 사례의 결과로 보아야 한다.")
    u=d["t2_under"]
    add_para(doc, f"Transformer v2도 클래스 0을 {u['majority_before']:,}건에서 {u['majority_after']:,}건으로 언더샘플링하여 학습 투입 건수를 {u['train_before']:,}건에서 {u['train_after']:,}건으로 축소했다.")
    add_heading(doc, "4. LLM 평가 데이터", 2)
    llm_rows = [
        ["LLM v1", "3,972", d["l1_met"]["evaluated_rows"], d["l1_met"]["skipped_missing_prediction"], pct(d["l1_met"]["evaluated_rows"]/3972), "캐시 0값 미반영 문제"],
        ["LLM v2", "3,972", d["l2_met"]["evaluated_rows"], d["l2_met"]["skipped_missing_prediction"], pct(d["l2_met"]["evaluated_rows"]/3972), "유효예측 행 평가"],
    ]
    add_caption(doc, "표 6. LLM 평가대상과 Coverage")
    add_table(doc, ["모형", "입력", "평가", "예측 결측", "Coverage", "비고"], llm_rows, [1200, 1000, 1000, 1100, 1100, 3960], 8.5)
    add_callout(doc, "비교 제한", "LLM은 전체자료 중 유효예측 행, Transformer는 독립 테스트 세트를 평가했다. 현 결과는 방향성 비교에는 사용할 수 있지만 통계적으로 공정한 모델 대결 결과는 아니다.", "yellow")

    add_heading(doc, "제4장. 모형 구축 및 평가 방법", 1)
    add_heading(doc, "1. LLM 분류", 2)
    add_para(doc, "LLM v1과 v2는 로컬 Ollama 환경의 gemma3:12b를 사용했다. 사업 메타데이터와 최대 16,000자의 사업설명자료 본문을 입력하고, BIOFIN 판정기준·경계규칙·대표사례를 포함한 시스템 프롬프트를 제공했다. 정답 레이블은 프롬프트에 포함하지 않고 분류 후 평가에만 사용했다. 결과는 label, confidence, reason, evidence 구조로 저장하며 캐시를 통해 중단 후 재개할 수 있도록 했다.")
    add_heading(doc, "2. Transformer 분류", 2)
    add_para(doc, "Transformer는 KLUE/RoBERTa-base로 긴 문서를 512토큰 단위로 나누고 인접 chunk 사이에 128토큰을 중첩했다. 각 chunk의 첫 토큰 임베딩을 학습 가능한 attention pooling으로 결합한 뒤 분류층에 입력했다. 기본 설정은 최대 5 epoch, 학습률 2×10⁻⁵, weight decay 0.01, warm-up 10%, dropout 0.1, attention hidden size 256, seed 42이다. 검증 Macro F1이 개선될 때 최적 모델을 저장하고 2 epoch 연속 개선이 없으면 조기 종료한다.")
    add_heading(doc, "3. 평가 지표", 2)
    metric_rows = [
        ["Accuracy", "전체 정답 건수 / 평가 건수", "전체 분류 성공률. 불균형에 민감"],
        ["Precision", "TP / (TP+FP)", "해당 클래스로 예측한 것의 정확성"],
        ["Recall", "TP / (TP+FN)", "실제 해당 클래스를 찾아낸 비율; 클래스별 정확도로 흔히 표현"],
        ["F1", "Precision·Recall 조화평균", "과대·과소 예측의 균형"],
        ["Macro F1", "클래스별 F1 단순평균", "소수 클래스를 동등하게 반영"],
        ["Weighted F1", "Support 가중 F1", "실제 분포를 반영하나 다수 클래스 영향 큼"],
        ["Coverage", "유효예측 / 전체입력", "미분류를 포함한 운영 완결성"],
    ]
    add_table(doc, ["지표", "개념", "본 보고서의 해석"], metric_rows, [1400, 2400, 5560], 8.3)

    page_break(doc)
    add_heading(doc, "제5장. 1차 카테고리 분류 결과", 1)
    add_heading(doc, "1. LLM v1", 2)
    m=d["l1_met"]
    add_para(doc, f"LLM v1의 기록상 평가대상은 {m['evaluated_rows']:,}건이며, {m['correct_rows']:,}건을 맞혀 Accuracy {pct(m['accuracy'])}, Macro Precision {fmt(m['macro_precision'])}, Macro Recall {fmt(m['macro_recall'])}, Macro F1 {fmt(m['macro_f1'])}을 보였다. 그러나 전체 3,972건 중 {m['skipped_missing_prediction']:,}건이 예측 결측으로 제외되었다.")
    add_callout(doc, "평가 파이프라인 이상", f"run_summary의 캐시 집계에는 클래스 0 예측이 1,062건, 빈 값이 5건 존재한다. 공식 평가에서는 이 합계와 같은 1,067건이 결측으로 제외되고 pred_0=0으로 집계됐다. 공식 평가모집단은 2,905건이며 이 안의 관련 사업은 503건이다. 원본 CSV의 0값을 유효예측으로 복구하면 3,967건(진짜 실패 5건 제외), 관련 사업 535건이 된다. 두 모집단은 이하에서 명확히 구분한다.", "yellow")
    l1rows=[]
    for code in sorted(m["per_class"], key=label_sort):
        x=m["per_class"][code]
        l1rows.append([code, UPPER_NAMES.get(code,""), x["support"], x["predicted"], fmt(x["precision"]), fmt(x["recall"]), fmt(x["f1"])])
    add_caption(doc, "표 7. LLM v1 클래스별 평가 결과")
    add_table(doc, ["코드", "명칭", "Support", "예측", "Precision", "Recall", "F1"], l1rows, [550, 3000, 900, 800, 1100, 1000, 2010], 7.7)
    add_para(doc, f"공식 유효 평가행에서는 클래스 1과 8의 F1이 각각 0.632, 0.564로 상대적으로 높았고, 클래스 2·4는 Recall이 0.776·0.805로 높지만 Precision이 약 0.08에 머물렀다. 공식 평가모집단의 관련 사업만 보면 정확 코드 일치율은 {pct(d['l1_official_nz']['accuracy'])}({d['l1_official_nz']['correct']}/{d['l1_official_nz']['n']})이다. 복구 모집단에서는 0으로 예측된 관련 사업 32건이 추가되어 {pct(d['l1_nz']['accuracy'])}({d['l1_nz']['correct']}/{d['l1_nz']['n']})가 된다.")
    add_heading(doc, "2. Transformer v1", 2)
    add_picture(doc, charts[2], "그림 3. Transformer 검증 성능 추이")
    add_para(doc, "Transformer v1은 epoch 2에서 검증 Macro F1 0.294로 최고점을 기록해 최적 모델로 선택되었다. 이후 학습 손실은 계속 감소했지만 검증 Macro F1은 낮아져 과적합 또는 불균형 자료에 대한 추가 적합 가능성이 나타났다.")
    tm=d["t1_met"]
    add_para(doc, f"최종 테스트 397건에서 Accuracy {pct(tm['test']['accuracy'])}, Macro F1 {fmt(tm['test']['macro_f1'])}, Weighted F1 {fmt(tm['test']['weighted_f1'])}이다. 관련 클래스만 대상으로 한 정확도는 {pct(d['t1_nz']['accuracy'])}({d['t1_nz']['correct']}/{d['t1_nz']['n']})이다.")
    t1rows=[]
    for code in sorted(d["t1_report"], key=label_sort):
        x=d["t1_report"][code]
        t1rows.append([code, UPPER_NAMES.get(code,""), x["support"], fmt(x["precision"]), fmt(x["recall"]), fmt(x["f1"])])
    add_caption(doc, "표 8. Transformer v1 테스트 클래스별 결과")
    add_table(doc, ["코드", "명칭", "Support", "Precision", "Recall", "F1"], t1rows, [600, 3300, 1000, 1100, 1100, 2260], 8)
    add_para(doc, "클래스 0은 Precision 0.985이나 Recall 0.584로, 비해당이라고 예측한 사례는 거의 맞지만 실제 비해당의 상당수를 관련 클래스로 보냈다. 클래스 2·9는 Recall 1.0, 클래스 6은 0.909였으나 Precision은 각각 0.146·0.290·0.333으로 낮아 관련 사업을 넓게 포착하는 대신 오탐이 많았다. 클래스 1·3·7·8은 테스트 표본이 1~3건이고 모두 Recall 0이었다.")
    add_heading(doc, "3. 1차 카테고리 모형 비교", 2)
    compare_rows = [
        ["평가 설계", "유효예측 2,905건", "독립 테스트 397건"],
        ["Accuracy", pct(d["l1_met"]["accuracy"]), pct(d["t1_met"]["test"]["accuracy"])],
        ["Macro F1", fmt(d["l1_met"]["macro_f1"]), fmt(d["t1_met"]["test"]["macro_f1"])],
        ["Coverage", pct(d["l1_met"]["evaluated_rows"]/3972), "100.0%"],
        ["주요 강점", "근거 문장·confidence 생성", "재현 가능한 홀드아웃 평가"],
        ["주요 한계", "0값 저장·평가 결함", "소수 클래스·과대포착"],
    ]
    add_table(doc, ["항목", "LLM v1", "Transformer v1"], compare_rows, [1700, 3830, 3830], 8.5)
    add_para(doc, "현 단계에서는 Transformer v1을 1차 분류의 기준 후보로 보는 것이 타당하다. 다만 두 모형이 동일 테스트자료에서 평가되지 않았고 LLM v1 파이프라인 결함이 있으므로, 최종 우위 판정은 0값 복구 후 공통 테스트 세트 재평가로 확정해야 한다.")

    page_break(doc)
    add_heading(doc, "제6장. 하위 카테고리 분류 결과", 1)
    add_heading(doc, "1. LLM v2", 2)
    m=d["l2_met"]
    zero_support=m["per_class"]["0"]["support"]
    add_para(doc, f"LLM v2는 전체 3,972건 중 {m['evaluated_rows']:,}건을 평가하여 Accuracy {pct(m['accuracy'])}, Macro Precision {fmt(m['macro_precision'])}, Macro Recall {fmt(m['macro_recall'])}, Macro F1 {fmt(m['macro_f1'])}을 기록했다. 평가대상의 {zero_support/m['evaluated_rows']:.1%}가 클래스 0이고, 0은 Precision·Recall·F1이 모두 1.0이었다.")
    add_callout(doc, "0/비0 지표 1.000의 의미", "유효예측 3,430건에서는 실제 0인 3,069건을 모두 0으로, 실제 비0인 361건을 모두 비0으로 예측해 경계 교차 오류가 없었다. 프롬프트에는 정답 상위·하위 레이블이 입력되지 않고 사업 메타데이터와 본문만 사용되므로 구조적으로 정답을 주입한 결과는 아니다. 다만 결측 542건에는 실제 0이 368건, 비0이 174건 포함되어 있다. 따라서 1.000은 선택된 유효예측 집합에 조건부인 결과이며, 전체자료 또는 독립 테스트에서 완벽하다는 뜻이 아니다. 결측 원인 점검과 공통 홀드아웃 재검증이 필요하다.", "yellow")
    add_para(doc, f"클래스 0을 제외한 관련 사업 {d['l2_nz']['n']}건의 정확한 하위 코드 일치율은 {pct(d['l2_nz']['accuracy'])}({d['l2_nz']['correct']}/{d['l2_nz']['n']})이다. 따라서 전체 Accuracy 95.36%는 비해당 판정에는 강하지만 관련 사업의 세부 분류까지 같은 수준으로 정확하다는 의미가 아니다.")
    good=[]; weak=[]
    for code,x in m["per_class"].items():
        if code=="0" or x["support"]==0: continue
        (good if x["f1"]>=.65 else weak if x["f1"]==0 else []).append((code,x))
    add_bullets(doc, [
        "상대적으로 우수: 2.03(F1 0.889), 2.04(0.948), 3.01(1.000), 5.06(0.700), 5.07(0.667), 6.01(0.696), 6.03(0.776), 7.04(0.667), 8.02(0.783).",
        "취약: 정답 표본이 있으나 F1이 0인 클래스가 다수이며, 특히 6.05는 Support 24인데 Recall 0으로 구조적 개선이 필요하다.",
        "과대예측: 5.03은 Recall 1.0이나 Precision 0.143, 7.01은 Recall 1.0이나 Precision 0.350으로 다른 사례를 해당 코드로 흡수한다.",
    ])
    add_heading(doc, "2. Transformer v2", 2)
    tm=d["t2_met"]
    add_para(doc, f"Transformer v2는 테스트 397건에서 Accuracy {pct(tm['test']['accuracy'])}, Macro F1 {fmt(tm['test']['macro_f1'])}, Weighted F1 {fmt(tm['test']['weighted_f1'])}을 기록했다. 관련 클래스 60건만 보면 정확한 하위 코드 일치율은 {pct(d['t2_nz']['accuracy'])}({d['t2_nz']['correct']}/{d['t2_nz']['n']})이다.")
    add_para(doc, "검증 loss는 metrics.json과 train_log.csv에서 NaN 또는 빈 값으로 기록되었다. 반면 Accuracy와 F1은 계산되어 모델 선택은 Macro F1 기준으로 수행됐다. 원인은 손실 계산 시 일부 클래스 가중치나 수치 연산 문제일 가능성이 있으나 산출물만으로 확정할 수 없으므로, 손실 함수 입력·class weight·mixed precision을 재현 실행으로 점검해야 한다.")
    t2rows=[]
    for code in sorted(d["t2_report"], key=label_sort):
        x=d["t2_report"][code]
        t2rows.append([code, SUB_NAMES.get(code,""), x["support"], fmt(x["precision"]), fmt(x["recall"]), fmt(x["f1"])])
    add_caption(doc, "표 9. Transformer v2 테스트 클래스별 결과")
    add_table(doc, ["코드", "명칭", "Support", "Precision", "Recall", "F1"], t2rows, [700, 3000, 900, 1050, 1000, 2710], 7.4)
    add_para(doc, "관련 클래스 가운데 9.05(F1 0.706), 7.01·8.02·9.02(각 0.667), 4.04(0.556), 6.03(0.500)이 상대적으로 높았다. 그러나 많은 클래스는 테스트 Support가 0~1건이며, 2.04는 Recall 0.813에도 Precision 0.134로 오탐이 많았다.")
    add_heading(doc, "3. 하위 카테고리 모형 비교", 2)
    add_table(doc, ["항목", "LLM v2", "Transformer v2"], [
        ["평가 설계", "유효예측 3,430건", "독립 테스트 397건"],
        ["Accuracy", pct(m["accuracy"]), pct(tm["test"]["accuracy"])],
        ["Macro F1", fmt(m["macro_f1"]), fmt(tm["test"]["macro_f1"])],
        ["관련 클래스 정확 일치율", pct(d["l2_nz"]["accuracy"]), pct(d["t2_nz"]["accuracy"])],
        ["Coverage", pct(m["evaluated_rows"]/3972), "100.0%"],
        ["특징", "0 식별 및 일부 빈번 클래스 강점", "일관된 일괄추론·확률 제공"],
    ], [1900, 3730, 3730], 8.3)
    add_para(doc, "원자료 기준 LLM v2의 관련 클래스 정확 일치율이 Transformer v2보다 높게 나타나지만, 모집단과 표본 구성이 다르므로 직접 우위로 확정할 수 없다. 공통 테스트 397건에 LLM v2를 다시 적용하여 McNemar 검정 또는 부트스트랩 차이 신뢰구간을 산출하는 것이 필요하다.")

    page_break(doc)
    add_heading(doc, "제7장. 네 모형 종합 비교", 1)
    add_heading(doc, "1. 클래스 0 포함·제외 결과", 2)
    nz_rows = [
        ["LLM v1", "1차", pct(d["l1_met"]["accuracy"]), f"{pct(d['l1_official_nz']['accuracy'])} ({d['l1_official_nz']['correct']}/{d['l1_official_nz']['n']})", "공식 평가 2,905건 기준"],
        ["Transformer v1", "1차", pct(d["t1_met"]["test"]["accuracy"]), f"{pct(d['t1_nz']['accuracy'])} ({d['t1_nz']['n']}건)", "관련 53건"],
        ["LLM v2", "하위", pct(d["l2_met"]["accuracy"]), f"{pct(d['l2_nz']['accuracy'])} ({d['l2_nz']['n']}건)", "0이 평가의 89.5%"],
        ["Transformer v2", "하위", pct(d["t2_met"]["test"]["accuracy"]), f"{pct(d['t2_nz']['accuracy'])} ({d['t2_nz']['n']}건)", "관련 60건"],
    ]
    add_table(doc, ["모형", "수준", "전체 Accuracy", "0 제외 정확 일치", "주의"], nz_rows, [1300, 900, 1500, 2100, 3560], 8.2)
    add_heading(doc, "2. 관련/비관련 이진 탐지 관점", 2)
    b_rows=[]
    for name,k,basis in [("LLM v1","l1_bin","0 복구 3,967건"),("Transformer v1","t1_bin","테스트 397건"),("LLM v2","l2_bin","유효예측 3,430건"),("Transformer v2","t2_bin","테스트 397건")]:
        b=d[k]
        b_rows.append([name, basis, b["tp"], b["fp"], b["fn"], b["tn"], fmt(b["precision"]), fmt(b["recall"]), fmt(b["f1"])])
    add_table(doc, ["모형", "계산 모집단", "TP", "FP", "FN", "TN", "Precision", "Recall", "F1"], b_rows, [1000,1700,600,600,600,600,1050,950,2260], 7.6)
    add_para(doc, "이진 지표는 ‘BIOFIN 관련 여부’를 먼저 선별하는 운영 시나리오를 보여준다. LLM v1만 공식 평가값으로 계산할 수 없어 CSV의 0 예측을 복구한 3,967건을 사용했으며, 다른 표의 공식 2,905건 지표와 직접 결합해서는 안 된다. LLM v2의 완벽한 이진 지표 역시 유효예측 집합에 조건부이다. Transformer는 정확한 세부 코드보다 관련 여부 탐지에서 더 나은 성능을 보여 계층형 운영의 근거가 된다.")
    add_heading(doc, "3. 오류 유형", 2)
    add_bullets(doc, [
        "비관련→관련 오탐: R&D·환경·기술 용어만으로 클래스 2나 4를 넓게 부여하는 유형.",
        "관련→비관련 누락: 관련성이 본문 일부에만 나타나거나 직접 목적이 약한 사업을 0으로 판단하는 유형.",
        "동일 상위범주 내 혼동: 연구(2.04), 교육(2.01/2.02), 인식(2.03)처럼 활동수단이 겹치는 경우.",
        "상위범주 간 혼동: 보호(7), 복원(8), 지속가능 이용(9)처럼 동일 생태공간에서 목적이 다른 경우.",
        "희소 클래스 미학습: 표본이 1~5건인 코드는 결정경계를 안정적으로 학습하거나 평가하기 어려움.",
    ])
    add_heading(doc, "4. 운영 특성", 2)
    add_table(doc, ["기준", "LLM", "Transformer"], [
        ["설명 가능성", "reason·evidence를 직접 생성", "attention·확률은 제공하나 자연어 근거는 별도"],
        ["재현성", "프롬프트·런타임 설정에 민감", "동일 체크포인트에서 상대적으로 높음"],
        ["신규 클래스", "프롬프트 수정으로 빠른 대응", "레이블 자료 확보와 재학습 필요"],
        ["추론 비용", "대형 모델 추론 및 긴 프롬프트 비용", "학습 후 대량 추론에 유리"],
        ["품질관리", "출력 형식·결측·환각 검증 필요", "분포 이동·confidence calibration 필요"],
    ], [1600, 3880, 3880], 8.5)

    add_heading(doc, "제8장. 결과 해석과 한계", 1)
    add_heading(doc, "1. 주요 해석", 2)
    add_para(doc, "첫째, 전체 Accuracy만으로는 실질 성능을 판단할 수 없다. LLM v2의 95.36%와 Transformer v2의 61.46%는 모두 클래스 0의 높은 비중에 영향을 받는다. 둘째, Macro F1이 0.225~0.309 수준이라는 점은 소수 클래스 전반의 균형 성능이 아직 낮음을 보여준다. 셋째, 관련 여부 탐지와 정확한 세부 코드 판정은 난도가 다르므로 운영도 두 단계로 나누는 것이 합리적이다.")
    add_heading(doc, "2. 평가상 한계", 2)
    add_bullets(doc, [
        "LLM과 Transformer가 동일한 테스트 세트에서 평가되지 않았다.",
        "LLM v1은 0 예측의 출력 또는 평가 처리 결함으로 현재 지표를 신뢰하기 어렵다.",
        "LLM v2는 예측 결측 542건이 제외되어 선택편향 가능성이 있다.",
        "하위 카테고리에는 검증·테스트에 정답 표본이 전혀 없는 클래스가 많다.",
        "Transformer v2의 검증 loss가 NaN으로 기록되어 학습 안정성 진단이 불완전하다.",
        "2023년 단일 연도 자료이므로 다른 연도·부처·정책 변화에 대한 일반화가 검증되지 않았다.",
    ])
    add_heading(doc, "3. 정답 레이블의 한계", 2)
    add_para(doc, "예산사업은 복수 목적과 활동을 동시에 가질 수 있으나 현재 문제는 하나의 정답 코드만 허용한다. 관련성이 간접적인 사업, 보호와 복원 활동이 함께 있는 사업, 정책·재정과 현장사업이 결합된 경우에는 단일 레이블이 실제 내용을 충분히 표현하지 못할 수 있다. 따라서 경계 사례를 재검수하고 다중 레이블 또는 주·부 카테고리 구조를 검토할 필요가 있다.")

    add_heading(doc, "제9장. 개선 방안 및 활용 제안", 1)
    add_heading(doc, "1. 우선순위별 개선과제", 2)
    improve_rows = [
        ["즉시", "LLM v1 0값 처리 수정·전수 재평가", "평가 신뢰성 복구", "필수"],
        ["즉시", "LLM·Transformer 공통 테스트 397건 구축", "공정 비교", "필수"],
        ["단기", "0 포함·제외, 이진 탐지, 계층 정확도 표준화", "운영지표 정합성", "높음"],
        ["단기", "Support 20 미만 클래스 추가 라벨링", "소수 클래스 안정화", "높음"],
        ["단기", "Transformer v2 NaN loss 재현·수정", "학습 안정성", "높음"],
        ["중기", "1차→하위 계층형 분류", "탐색공간 축소", "높음"],
        ["중기", "LLM·Transformer 합의/불일치 라우팅", "검토 효율", "중간"],
        ["장기", "다년도 외부검증과 지속 재학습", "일반화·유지관리", "높음"],
    ]
    add_table(doc, ["시점", "과제", "기대효과", "우선도"], improve_rows, [900, 4000, 2900, 1560], 8.2)
    add_heading(doc, "2. 권고 운영 절차", 2)
    add_bullets(doc, [
        "1단계 관련 여부 판정: 클래스 0 대 비0을 높은 Recall 목표로 선별한다.",
        "2단계 상위 카테고리 판정: Transformer v1과 LLM의 합의 여부를 확인한다.",
        "3단계 하위 카테고리 판정: 상위 범주 내 후보만 비교하고 확률·근거를 함께 저장한다.",
        "4단계 사람 검토: 낮은 confidence, 모형 불일치, 희소 클래스, 복수 목적 사업을 우선 검토한다.",
        "5단계 환류: 확정 결과를 학습자료에 누적하고 정기적으로 성능·분포 이동을 점검한다.",
    ])
    add_heading(doc, "3. 자동 확정 기준 제안", 2)
    add_para(doc, "현 시점에서는 전체 자동 확정보다 위험 기반 반자동화가 적절하다. 클래스별 검증 표본이 충분하고 최근 공통 테스트에서 Precision과 Recall 기준을 모두 충족하며 두 모형이 동일 코드에 합의한 경우만 자동 확정 후보로 설정한다. 그 외는 사람 검토 대상으로 보낸다. 임계값은 비용함수와 정책적 오류 허용도에 따라 별도 검증자료에서 결정해야 한다.")

    add_heading(doc, "제10장. 결론", 1)
    add_para(doc, "본 연구는 3,972건의 2023년 예산사업을 대상으로 LLM과 Transformer를 활용한 BIOFIN 자동분류 가능성을 검토했다. Transformer v1은 독립 테스트에서 60.45%의 1차 분류 정확도를 보였으며, LLM v2는 유효예측 자료에서 95.36%의 하위 분류 정확도를 기록했다. 그러나 클래스 0 편중, 희소 클래스, 평가 모집단 차이와 LLM v1 처리 오류 때문에 단일 수치만으로 모형을 선정해서는 안 된다.")
    add_para(doc, "실무적으로는 관련 여부→상위→하위의 계층형 구조와 사람 검토를 결합하는 것이 가장 현실적이다. 우선 LLM v1 평가 파이프라인을 수정하고 공통 테스트 세트에서 네 모형을 재평가해야 한다. 이후 클래스별 충분한 표본, 신뢰도 보정, 다년도 외부검증을 확보하면 자동 확정 범위를 단계적으로 확대할 수 있다.")

    page_break(doc)
    add_heading(doc, "부록 1. 하위 카테고리 코드", 1)
    subrows=[[k,v] for k,v in sorted(SUB_NAMES.items(), key=lambda x: label_sort(x[0]))]
    add_table(doc, ["코드", "명칭"], subrows, [1300,8060], 8.3)

    add_heading(doc, "부록 2. 전체 레이블 분포", 1)
    dist1={r["label"]:int(r["count"]) for r in d["t1_dist"]}
    dist2={r["label"]:int(r["count"]) for r in d["t2_dist"]}
    add_heading(doc, "1. 1차 카테고리", 2)
    add_table(doc, ["코드", "명칭", "전체 건수", "비율"], [[k,UPPER_NAMES[k],dist1.get(k,0),pct(dist1.get(k,0)/3972,2)] for k in UPPER_NAMES], [700,4300,1500,2860], 8.2)
    add_heading(doc, "2. 하위 학습 클래스", 2)
    label_map=read_json(T2/"label_map.json")["code_to_id"]
    add_table(doc, ["코드", "명칭", "전체 건수", "비율"], [[k,SUB_NAMES.get(k,""),dist2.get(str(label_map[k]),0),pct(dist2.get(str(label_map[k]),0)/3972,2)] for k in sorted(label_map,key=label_sort)], [800,4000,1500,3060], 7.8)
    add_para(doc, "주: Transformer v2의 label_distribution.csv는 내부 class ID로 저장되어 label_map.json을 이용해 원래 하위 코드로 변환하였다.", italic=True, color=MUTED)

    add_heading(doc, "부록 3. 산출물 및 재현 설정", 1)
    add_table(doc, ["모형", "핵심 산출물", "용도"], [
        ["LLM v1", "run_summary.json, evaluation_metrics.json, confusion_matrix.csv, incorrect_predictions.csv", "실행·전체평가·오분류"],
        ["LLM v2", "run_summary.json, evaluation_metrics.json, review_needed.csv, subcategory_label_audit.csv", "하위평가·검토·감사"],
        ["Transformer v1", "best_model.pt, metrics.json, classification_report.txt, split_summary.json, test_predictions.csv", "모델·평가·분할"],
        ["Transformer v2", "best_model.pt, label_map.json, metrics.json, classification_report.txt, test_predictions.csv", "하위 모델·코드변환·평가"],
    ], [1400, 5000, 2960], 8)
    add_heading(doc, "자료 출처", 2)
    add_bullets(doc, [
        str(L1), str(L2), str(T1), str(T2),
        str(ROOT / "llm/v1/classify_biofin_category_with_ollama.py"),
        str(ROOT / "llm/v2/classify_biofin_subcategory_with_ollama.py"),
        str(ROOT / "transformer/v1/src/train_attention_classifier.py"),
        str(ROOT / "transformer/v2/src/train_attention_classifier.py"),
    ])
    add_para(doc, "경로 확인: Transformer v2 분석대상은 사용자가 지정한 2026-08-18 실행 폴더 transformer/v2/outputs/260818_2023data이다. 실제로 best_model.pt, metrics.json, classification_report.txt, label_map.json과 test_predictions.csv가 존재한다. 과거 실행 폴더인 transformer/v2/outputs/260812_subcategory_v3는 이번 네 모형 비교 범위에 포함하지 않았다.", italic=True, color=MUTED)
    add_para(doc, "본 보고서의 추가 산출 지표(클래스 0 제외 정확 일치율 및 관련/비관련 이진 지표)는 각 classified/test_predictions CSV의 정답·예측 열을 기준으로 계산하였다.", italic=True, color=MUTED)

    # Keep table rows from splitting when small and set all fonts.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Malgun Gothic"
                        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name is None:
                run.font.name = "Malgun Gothic"
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.core_properties.title = "BIOFIN 예산사업 분류모형 개발 결과 및 비교·평가 보고서"
    doc.core_properties.subject = "LLM·Transformer 기반 BIOFIN 예산분류 결과"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "BIOFIN, 예산사업, LLM, Transformer, 분류모형"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
